"""
Data Analysis and Statistical Utility

This script provides tools for analyzing datasets, calculating statistical metrics, and generating detailed reports. It includes functionality for descriptive statistics, hypothesis testing, and inequality indices, as well as tools for formatting results into a readable document format (e.g., DOCX).

Key functionalities include:
- Generating descriptive statistics for categorical and numerical data.
- Performing statistical tests (e.g., chi-square, t-tests, ANOVA, Mann-Whitney U).
- Calculating inequality measures like Gini and Theil indices.
- Exporting analysis results to Word documents.
- Handling large datasets with bootstrapping techniques.

Dependencies:
- pandas
- numpy
- scipy
- statsmodels
- sklearn
- docx
- tqdm

Usage:
1. Import the desired functions into your analysis script or Jupyter Notebook.
2. Use `describe()` to generate descriptive statistics.
3. Perform hypothesis testing with built-in functions.
4. Save results using the `add_to_docx()` function for presentation-ready tables.

Example:
    from analyse import describe, add_to_docx

    # Sample data
    df = pd.DataFrame({
        'group': ['A', 'A', 'B', 'B', 'C', 'C'],
        'value': [1.1, 2.3, 1.5, 2.7, 3.1, 2.9]
    })

    describe(df, factors=['value'], group_var='group', table_name='Descriptive Analysis')

    add_to_docx(results, table_name="Analysis Results", output_dir="./output")

Authors: [Your Name Here]
Version: 1.0
License: MIT License
"""

# Standard library imports
import os

# Third-party library imports
import pandas as pd
import numpy as np
from pandas.api.types import is_numeric_dtype
from sklearn.utils import resample
from scipy import stats
from scipy.stats import shapiro, chi2_contingency, anderson
from tqdm import tqdm
from tabulate import tabulate
from docx import Document

class ResultExport:
    """
    A class for export the result to docx.
    """
    
    @staticmethod
    def add_to_docx(results, table_name, output_dir):
        """
        Export analysis results to a Word document.

        Parameters:
        - results (DataFrame or list): Analysis results to include in the document.
        - table_name (str): Title of the table in the document.
        - output_dir (str): Directory to save the document.
        
        Returns:
        - None
        """
        
        doc = Document()
        
        # Add a title for the table
        doc.add_heading(table_name, level=2)
        
        # Check if results is a DataFrame or list and if it's empty
        if isinstance(results, pd.DataFrame):
            if results.empty:
                doc.add_paragraph("No data available.")
            else:
                # Convert DataFrame to list of dictionaries
                results = results.to_dict(orient='records')
        elif not results:
            doc.add_paragraph("No data available.")
            return

        # Add a table with a header row if there is data
        if results:
            headers = results[0].keys()
            table = doc.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            # Add the data rows
            for row_data in results:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(row_data[header])
        
        # Save the document
        doc.save(f'{output_dir}/{table_name}.docx')

class Descriptive:
    """
    A class for descriptive analysis
    - Provides methods for descriptive analysis of categorical and numerical variables.
    """
    
    @staticmethod
    def des_cat(df_input, factor, group_var=None, p_value=None):
        """
        Generate descriptive statistics for a categorical variable.

        This method calculates the total and group-level distribution for a categorical
        variable (`factor`). If a grouping variable (`group_var`) is provided, the analysis
        includes group-level breakdowns.

        Parameters:
        - df_input (DataFrame): The input dataset.
        - factor (str): The column name of the categorical variable to analyze.
        - group_var (str, optional): The column name of the grouping variable for stratified analysis.
        - p_value (str, optional): The p-value from a statistical test, if available.

        Returns:
        - list of dict: A list of dictionaries representing the descriptive statistics
          for the categorical variable. Each dictionary corresponds to a row in the results.
        """
        
        rows = []
        
        if group_var:
            df = df_input.dropna(subset=[group_var])

        # Header row
        header_row = {'Characteristics': factor, 'Total': '', 'p-value': p_value}
        if group_var:
            groups = df[group_var].dropna().unique()
            for group in groups:
                header_row[group] = ''
        rows.append(header_row)

        # Calculate total for each category
        categories = df[factor].dropna().unique()
        grand_total = len(df[factor].dropna())

        for category in categories:
            subtotal = len(df[df[factor] == category])
            percent = (subtotal / grand_total) * 100
            row = {'Characteristics': category, 'Total': f'{subtotal:,}\n({percent:.2f}%)'}

            if group_var:
                for group in groups:
                    group_total = len(df[(df[group_var] == group) & (df[factor].notna())])
                    group_count = len(df[(df[factor] == category) & (df[group_var] == group)])
                    group_percent = (group_count / group_total) * 100
                    row[group] = f'{group_count:,}\n({group_percent:.2f}%)'

            row['p-value'] = ''  # Add p-value column placeholder
            rows.append(row)

        return rows

    @staticmethod
    def des_num(df, factor, group_var=None, p_value=None, is_normal=True):
        """
        Generate descriptive statistics for a numerical variable.

        This method calculates summary statistics (mean and standard deviation
        or median and interquartile range) for a numerical variable (`factor`).
        It can provide group-level breakdowns if a grouping variable (`group_var`)
        is specified.

        Parameters:
        - df (DataFrame): The input dataset.
        - factor (str): The column name of the numerical variable to analyze.
        - group_var (str, optional): The column name of the grouping variable for stratified analysis.
        - p_value (str, optional): The p-value from a statistical test, if available.
        - is_normal (bool, optional): Specifies whether the data is normally distributed.
          - If True, calculates mean and standard deviation.
          - If False, calculates median and interquartile range.

        Returns:
        - dict: A dictionary containing the descriptive statistics for the numerical variable.
          - Includes the overall summary and group-level summaries if `group_var` is provided.
        """
        
        rows = {}

        if is_normal:
            mean = df[factor].dropna().mean()
            sd = df[factor].dropna().std()
            rows = {
                'Characteristics': factor,
                'Total': f'{mean:.2f}\n({sd:.2f})',
                'p-value': p_value
            }

            if group_var:
                groups = df[group_var].dropna().unique()
                for group in groups:
                    group_df = df[df[group_var] == group]
                    submean = group_df[factor].dropna().mean()
                    subsd = group_df[factor].dropna().std()
                    rows[group] = f'{submean:.2f}\n({subsd:.2f})'

        else:
            median = df[factor].dropna().median()
            p25 = df[factor].dropna().quantile(0.25)
            p75 = df[factor].dropna().quantile(0.75)
            rows = {
                'Characteristics': factor,
                'Total': f'{median:.2f}\n({p25:.2f}-{p75:.2f})',
                'p-value': p_value
            }

            if group_var:
                groups = df[group_var].dropna().unique()
                for group in groups:
                    group_df = df[df[group_var] == group]
                    submedian = group_df[factor].dropna().median()
                    subp25 = group_df[factor].dropna().quantile(0.25)
                    subp75 = group_df[factor].dropna().quantile(0.75)
                    rows[group] = f'{submedian:.2f}\n({subp25:.2f}-{subp75:.2f})'

        return rows

    @staticmethod
    def show_median(df, column):
        """
        Calculate the median and interquartile range for a numerical column.

        This function computes the median, 25th percentile (Q1), and 75th percentile (Q3)
        for a specified numerical column and returns them in a formatted string.

        Parameters:
        - df (DataFrame): The input dataset.
        - column (str): The name of the numerical column to analyze.

        Returns:
        - str: A string containing the median and interquartile range in the format:
        `median (Q1 - Q3)`.
        """
        
        median = np.median(df[column])
        p25 = np.percentile(df[column], 25)
        p75 = np.percentile(df[column], 75)

        result = f'{median:.2f} ({p25:.2f} - {p75:.2f})'

        return result

    @staticmethod
    def describe(df, factors, group_var, overall=False, bootstrap=False, n_bootstraps=1000, table_name=None):
        """
        Generate descriptive statistics for categorical and numerical variables.

        This function calculates descriptive statistics for the specified columns (`factors`),
        grouped by a specified variable (`group_var`). It provides both overall and group-level
        summaries and can optionally perform bootstrap-based significance testing.

        Parameters:
        - df (DataFrame): The input dataset.
        - factors (list): List of column names to analyze (categorical or numerical).
        - group_var (str): The name of the grouping variable for stratified analysis.
        - overall (bool, optional): Whether to include overall statistics in the output. Defaults to False.
        - bootstrap (bool, optional): Whether to use bootstrapping for p-value calculations. Defaults to False.
        - n_bootstraps (int, optional): Number of bootstrap iterations to perform. Defaults to 1000.
        - table_name (str, optional): The title of the output table.

        Returns:
        - None: The results are printed to the console and saved as a Word document.

        Notes:
        - Categorical variables are analyzed using chi-square tests.
        - Numerical variables are analyzed using t-tests (for two groups) or ANOVA (for more than two groups).
        - The results include overall statistics and group-level breakdowns for each factor.
        - The output is formatted as a table and saved as a Word document using `add_to_docx()`.
        """
        
        print(f"{table_name}")
        
        results = pd.DataFrame(columns=["Characteristics", "Total"] + list(df[group_var].dropna().unique()) + ["p-value"])

        for factor in factors:

            # Descriptive analysis for categorical variables
            if df[factor].dtype == 'O':
                # Chi-square with Bootstrapping and Downsampling
                contingency_table = pd.crosstab(df[factor], df[group_var])
                group_sizes = contingency_table.sum(axis=0)
                min_group_size = group_sizes.min()
                
                if bootstrap:
                    p_values = []
                    for _ in tqdm(range(n_bootstraps), desc=f"Bootstrapping {factor}", leave=False):
                        # Downsample each group to the size of the smallest group
                        sampled_df = pd.concat([df[df[group_var] == group].sample(min_group_size, replace=True) 
                                                for group in df[group_var].dropna().unique()])
                        sampled_contingency = pd.crosstab(sampled_df[factor], sampled_df[group_var])
                        _, p_val, _, _ = chi2_contingency(sampled_contingency)
                        p_values.append(p_val)
                    p_value_mean = np.mean(p_values)
                    ci_lower, ci_upper = np.percentile(p_values, [2.5, 97.5])
                    p_value = f"{p_value_mean:.2f}\n(95% CI: {ci_lower:.2f}-{ci_upper:.2f})"
                else:
                    contingency_table = pd.crosstab(df[factor], df[group_var])
                    _, p_value, _, _ = chi2_contingency(contingency_table)
                    p_value = f"{p_value:.2f}"

                des = Descriptive.des_cat(df, factor, group_var, p_value)
                descriptive_df = pd.DataFrame(des)

            # Descriptive analysis for numerical variables
            elif is_numeric_dtype(df[factor]):
                
                # Normality test before conducting the descriptive analysis
                normality_pvalues = {}
                for group in df[group_var].dropna().unique():
                    group_data = df[df[group_var] == group][factor].dropna()
                    if len(group_data) > 5000:
                        ad_stat, critical_values, _ = anderson(group_data)
                        pvalue = (ad_stat > critical_values[-1])
                        normality_pvalues[group] = 0 if pvalue else 1
                    else:
                        if len(group_data) > 3:
                            stat, pvalue = shapiro(group_data)
                            normality_pvalues[group] = pvalue
                
                is_normal = all(p > 0.05 for p in normality_pvalues.values())
                groups = [df[df[group_var] == group][factor].dropna() for group in df[group_var].dropna().unique()]
                min_group_size = min([len(group) for group in groups])

                # For normal distribution
                if is_normal:
                    
                    # 2 groups : T-Test
                    if len(groups) == 2 and all(len(group) > 1 for group in groups):
                        if bootstrap:
                            p_values = []
                            for _ in tqdm(range(n_bootstraps), desc=f"Bootstrapping {factor}", leave=False):
                                sample_group1 = resample(groups[0], n_samples=min_group_size, replace=True)
                                sample_group2 = resample(groups[1], n_samples=min_group_size, replace=True)
                                _, p_val = stats.ttest_ind(sample_group1, sample_group2, equal_var=False)
                                p_values.append(p_val)
                            p_value_mean = np.mean(p_values)
                            ci_lower, ci_upper = np.percentile(p_values, [2.5, 97.5])
                            p_value = f"{p_value_mean:.2f}\n(95% CI: {ci_lower:.2f}-{ci_upper:.2f})"
                        else:
                            sample_group1 = resample(groups[0], n_samples=min_group_size, replace=False)
                            sample_group2 = resample(groups[1], n_samples=min_group_size, replace=False)
                            _, p_value = stats.ttest_ind(sample_group1, sample_group2, equal_var=False)
                            p_value = f"{p_value:.2f}"
                        print(f'Testing for {factor} using T-test')
                    
                    # > 2 groups : ANOVA
                    elif len(groups) > 2 and all(len(group) > 1 for group in groups):
                        if bootstrap:
                            p_values = []
                            for _ in tqdm(range(n_bootstraps), desc=f"Bootstrapping {factor}", leave=False):
                                sampled_groups = [resample(group, n_samples=min_group_size, replace=True) for group in groups]
                                _, p_val = stats.f_oneway(*sampled_groups)
                                p_values.append(p_val)
                            p_value_mean = np.mean(p_values)
                            ci_lower, ci_upper = np.percentile(p_values, [2.5, 97.5])
                            p_value = f"{p_value_mean:.2f}\n(95% CI: {ci_lower:.2f}-{ci_upper:.2f})"
                        else:
                            sampled_groups = [resample(group, n_samples=min_group_size, replace=False) for group in groups]
                            _, p_value = stats.f_oneway(*sampled_groups)
                            p_value = f"{p_value_mean:.2f}"
                        print(f'Testing for {factor} using ANOVA')
                    else:
                        p_value = float('nan')
                        
                # For non-normal distribution        
                else:
                    
                    # 2 groups : Mann-Whitney U test
                    if len(groups) == 2 and all(len(group) > 1 for group in groups):
                        if bootstrap:
                            p_values = []
                            for _ in tqdm(range(n_bootstraps), desc=f"Bootstrapping {factor}", leave=False):
                                sample_group1 = resample(groups[0], n_samples=min_group_size, replace=True)
                                sample_group2 = resample(groups[1], n_samples=min_group_size, replace=True)
                                _, p_val = stats.mannwhitneyu(sample_group1, sample_group2)
                                p_values.append(p_val)
                            p_value_mean = np.mean(p_values)
                            ci_lower, ci_upper = np.percentile(p_values, [2.5, 97.5])
                            p_value = f"{p_value_mean:.2f}\n(95% CI: {ci_lower:.2f}-{ci_upper:.2f})"
                        else:
                            sample_group1 = resample(groups[0], n_samples=min_group_size, replace=False)
                            sample_group2 = resample(groups[1], n_samples=min_group_size, replace=False)
                            _, p_value = stats.mannwhitneyu(sample_group1, sample_group2)
                            p_value = f"{p_value:.2f}"
                        print(f'Testing for {factor} using Mann-Whitney U')
                    
                    # > 2 groups : Kruskal-Wallis H Test
                    elif len(groups) > 2 and all(len(group) > 1 for group in groups):
                        if bootstrap:
                            p_values = []
                            for _ in tqdm(range(n_bootstraps), desc=f"Bootstrapping {factor}", leave=False):
                                sampled_groups = [resample(group, n_samples=min_group_size, replace=True) for group in groups]
                                _, p_val = stats.kruskal(*sampled_groups)
                                p_values.append(p_val)
                            p_value_mean = np.mean(p_values)
                            ci_lower, ci_upper = np.percentile(p_values, [2.5, 97.5])
                            p_value = f"{p_value_mean:.2f}\n(95% CI: {ci_lower:.2f}-{ci_upper:.2f})"
                        else:
                            _, p_value = stats.kruskal(*groups)
                            p_value = f"{p_value:.2f}"
                        print(f'Testing for {factor} using Kruskal-Wallis H')
                    else:
                        p_value = f"N/A"

                des = Descriptive.des_num(df, factor, group_var, p_value, is_normal)
                descriptive_df = pd.DataFrame([des])

            # Adding the new results
            results = pd.concat([results, descriptive_df], ignore_index=True)

        # Whether want to show the overall columns
        if not overall:
            results.drop(columns=['Total'], inplace=True)

        # Print the results and saved the docx
        print(tabulate(results, showindex=False, headers="keys"))
        ResultExport.add_to_docx(results, table_name, output_dir='output/analyse')

def coverage(df_raw, value_column, level):
    """
    Calculate coverage for a specified level of geographic or organizational granularity.

    This function computes the coverage percentage for a specific level (e.g., region, province, or district)
    based on the count of non-zero entries in a given column.

    Parameters:
    - df_raw (DataFrame): The input dataset.
    - value_column (str): The column to check for non-zero values (e.g., service availability).
    - level (str): The granularity level for the calculation. Options are:
      - 'hregion': Health region (13 units).
      - 'prov': Province (77 units).
      - 'dist': District (928 units).

    Returns:
    - str: A string in the format: `n (percentage%)`, where `n` is the count of non-zero entries
      and `percentage` is the coverage relative to the total units at the specified level.

    Notes:
    - The `level` parameter maps to predefined total counts:
      - 'hregion': 13
      - 'prov': 77
      - 'dist': 928
    - Raises a `ValueError` if the specified `level` is invalid.
    - Useful for analyzing service coverage or availability across predefined units.
    """

    # Create a copy of the input dataframe to avoid modifying the original
    df = df_raw.copy()

    # Count the non-zero values in the specified value column
    n = len(df[df[value_column] != 0])

    # Define the map for levels and corresponding total counts
    map_N = {
        'hregion': 13,
        'prov': 77,
        'dist': 928
    }

    # Ensure that the level exists in the map
    if level not in map_N:
        raise ValueError(f"Invalid level: {level}. Choose from {list(map_N.keys())}")

    # Calculate the percentage
    percent = (n / map_N[level]) * 100

    result = f'{n} ({percent:.2f}%)'
    
    return result

def theil_T(df_raw, h_column, pop_column):
    """
    Calculate Theil's T Index for inequality measurement.

    This function computes Theil's T index, which quantifies inequality in the
    distribution of a resource (e.g., income, healthcare services) relative to the
    population.

    Parameters:
    - df_raw (DataFrame): The input dataset containing resource and population data.
    - h_column (str): The column name representing the resource variable (e.g., income, healthcare units).
    - pop_column (str): The column name representing the population variable.

    Returns:
    - float: The calculated Theil's T index.

    Notes:
    - Theil's T index is calculated using the formula:
      T = Σ [(h_i / H) * log((h_i / H) / (p_i / P))]
      where h_i and p_i are the resource and population values for each unit, and
      H and P are the totals for the resource and population, respectively.
    - The index ranges from 0 (perfect equality) to higher values indicating greater inequality.
    - Rows with zero resource values are excluded from the calculation.
    """
    
    df = df_raw.copy()
    H = np.sum(df[h_column])
    P = np.sum(df[pop_column])
    df['t'] = df.apply(lambda row: (row[h_column] / H) * np.log((row[h_column] / H) / (row[pop_column] / P)) if row[h_column] > 0 else 0, axis=1)
    T = np.sum(df['t'])
    
    return T

def gini(df_raw, h_column, pop_column):
    """
    Calculate the Gini Coefficient for inequality measurement.

    This function computes the Gini coefficient, a widely used measure of inequality in
    resource distribution, based on a resource variable (e.g., income, healthcare units)
    and a population variable.

    Parameters:
    - df_raw (DataFrame): The input dataset containing resource and population data.
    - h_column (str): The column name representing the resource variable.
    - pop_column (str): The column name representing the population variable.

    Returns:
    - float: The calculated Gini coefficient.

    Notes:
    - The Gini coefficient ranges from 0 (perfect equality) to 1 (maximum inequality).
    - The formula involves sorting the data by the ratio of resource to population,
      calculating cumulative distributions, and summing weighted differences.
    - The calculation uses a stepwise approach:
      - Calculate cumulative proportions of the population and resources.
      - Compute the area under the Lorenz curve.
      - Derive the Gini coefficient as 1 minus the normalized area under the Lorenz curve.
    """
    
    df = df_raw.copy()
    H = np.sum(df[h_column])
    P = np.sum(df[pop_column])
    
    df['x'] = df.apply(lambda row: (row[h_column] / row[pop_column]), axis=1)
    df = df.sort_values(by='x').reset_index(drop=True)

    df['p'] = df.apply(lambda row: (row[pop_column] / P), axis=1)
    df['h'] = df.apply(lambda row: (row[h_column] / H), axis=1)
    
    df['F'] = df['p'].cumsum()
    df['F-1'] = df['F'] - df['p']
    df['R'] = df['h'].cumsum()
    df['R-1'] = df['R'] - df['h']

    df['g'] = (df['F'] - df['F-1']) * (df['R'] + df['R-1'])

    G = 1 - (np.sum(df['g']))

    return G

